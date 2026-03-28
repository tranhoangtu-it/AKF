"""
AKF + python-docx — AI provenance metadata in Word documents

Embed trust scores and provenance into DOCX files via custom XML parts.
When someone opens the Word doc, the provenance travels with it.

Usage:
    pip install akf python-docx
    python python_docx_provenance.py

Learn more: https://akf.dev
"""

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from akf import stamp
import json
import tempfile
import os


def stamp_docx_with_akf(docx_path, agent="claude", evidence="generated from data"):
    """Stamp a DOCX file with AKF trust metadata via custom properties."""

    # Create AKF trust metadata
    unit = stamp(
        content=f"AI-generated document: {os.path.basename(docx_path)}",
        confidence=0.85,
        agent=agent,
        evidence=[evidence],
    )

    # Use AKF's built-in DOCX embedding
    from akf import universal as akf_u
    akf_u.embed(docx_path, metadata=unit.to_dict(compact=True))

    return unit


def read_akf_from_docx(docx_path):
    """Read AKF trust metadata from a DOCX file."""
    from akf import universal as akf_u
    return akf_u.extract(docx_path)


if __name__ == "__main__":
    print("=== AKF + python-docx: AI Provenance in Word Documents ===\n")

    # Create a sample DOCX
    doc = Document()
    doc.add_heading("Q3 Revenue Report", 0)
    doc.add_paragraph("Revenue was $4.2B, up 12% YoY.")
    doc.add_paragraph("Cloud segment grew 15%.")

    docx_path = tempfile.mktemp(suffix=".docx")
    doc.save(docx_path)
    print(f"1. Created DOCX: {docx_path}")

    # Stamp with AKF
    unit = stamp_docx_with_akf(docx_path, agent="claude-3.5", evidence="SEC 10-Q analysis")
    print(f"2. Stamped with AKF trust metadata")
    print(f"   Trust: {unit.claims[0].confidence}")
    print(f"   Agent: {unit.agent}")

    # Read it back
    akf_data = read_akf_from_docx(docx_path)
    if akf_data:
        print(f"\n3. Read back from DOCX:")
        print(f"   {json.dumps(akf_data, indent=2)[:200]}...")

    # Cleanup
    os.unlink(docx_path)
    print(f"\n✅ Trust metadata embedded in DOCX custom XML")
    print(f"   Open in Word → File → Properties to see AKF metadata")
    print(f"   Learn more: https://akf.dev")
